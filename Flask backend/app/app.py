import sys
import time
import redis
import base64
import secrets
import werkzeug
from PIL import Image
import werkzeug.exceptions
from flask import Flask, request, Response
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from io import BytesIO
import docx
from docx.shared import Inches, Cm


class ErrorCode:
    IncorrectUserToken = 0
    IncorrectAdminToken = 1
    StringTooLong = 2
    StringFailedFiltering = 3
    InternalError = 4
    ImageWidthInvalid = 5
    ImageHeightInvalid = 6
    InvalidImageIndex = 7
    InvalidCommentIndex = 8
    IncorrectAlbumId = 9
    InvalidAlbumIndex = 10


def success(value):
    return {"success": True, "params": value}, 200


def error(code: int):
    return {"success": False, "params": code}, 200


def image_id_list_name(user_token: str):
    return f"{user_token}_iid"


def image_stream_name(user_token: str):
    return f"{user_token}_i"


def image_comment_id_list_name(user_token: str, image_id: str):
    return f"{user_token}_{image_id}_cid"


def image_comment_stream_name(user_token: str, image_id: str):
    return f"{user_token}_{image_id}_c"


def album_id_list_name(user_token: str):
    return f"{user_token}_aid"


def album_stream_name(user_token: str):
    return f"{user_token}_a"


def album_image_id_list_name(user_token: str, album_id: str):
    return f"{user_token}_{album_id}_iid"


#def album_image_comment_stream_name(user_token: str,album_id: str):
#    return f"{user_token}_{album_id}_i"

#@TODO: User provided strings should be filtered in some way to at least prevent basic swears or bad phrases.
def filter_user_string(string: str) -> bool:
    return True


#@TODO: This should log to a file.
def log_endpoint_error(string: str):
    print(string, file=sys.stderr)


app = Flask(__name__)
database = redis.StrictRedis(host="redis", port=6379, decode_responses=True)


def does_event_exist(user_token: str) -> bool:
    return database.sismember("user_tokens", user_token) == 1


def is_admin_token_valid(user_token: str, admin_token: str) -> bool:
    return database.hget("admin_tokens", user_token) == admin_token


@app.errorhandler(werkzeug.exceptions.InternalServerError)
def handle_internal_server_error(error):
    log_endpoint_error(str(error))
    return {"success": False, "params": ErrorCode.InternalError}, 500


@app.errorhandler(werkzeug.exceptions.NotFound)
def handle_not_found_error(error):
    log_endpoint_error(str(error))
    return {"success": False, "params": ErrorCode.InternalError}, 404


################################################################
#                     ENDPOINTS (EVENTS)                       #
################################################################

@app.post("/v0/event")
def endpoint_create_event():
    if "event_name" not in request.json:
        return error(ErrorCode.InternalError)
    event_name = str(request.json["event_name"])
    if len(event_name) > 128:
        return error(ErrorCode.StringTooLong)
    if not filter_user_string(event_name):
        return error(ErrorCode.StringFailedFiltering)

    user_token = "event_" + str(database.incr("user_token_incr"))
    admin_token = secrets.token_hex(3)

    transaction = database.pipeline()
    transaction.sadd("user_tokens", user_token)
    transaction.hset("admin_tokens", user_token, admin_token)
    transaction.hset("event_names", user_token, event_name)
    transaction.execute()
    return success({"user_token": user_token, "admin_token": admin_token})


@app.delete("/v0/event/<user_token>")
def endpoint_delete_event(user_token: str):
    if "admin_token" not in request.json:
        return error(ErrorCode.InternalError)
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    admin_token = str(request.json["admin_token"])
    if not is_admin_token_valid(user_token, admin_token):
        return error(ErrorCode.IncorrectAdminToken)

    image_stream = database.xrange(image_stream_name(user_token), "0-0", "+")
    album_stream = database.xrange(album_stream_name(user_token), "0-0", "+")
    transaction = database.pipeline()
    for image_id, _ in image_stream:
        transaction.ltrim(image_comment_id_list_name(user_token, image_id), 0, -1)
        transaction.xtrim(image_comment_stream_name(user_token, image_id), minid="0-0")
    for album_id, _ in album_stream:
        transaction.ltrim(album_image_id_list_name(user_token, album_id), 0, -1)
    transaction.ltrim(album_id_list_name(user_token), 0, -1)
    transaction.xtrim(album_stream_name(user_token), minid="0-0")
    transaction.ltrim(image_id_list_name(user_token), 0, -1)
    transaction.xtrim(image_stream_name(user_token), minid="0-0")
    transaction.hdel("event_names", user_token)
    transaction.hdel("admin_tokens", user_token)
    transaction.srem("user_tokens", user_token)
    transaction.execute()
    return success({})


@app.get("/v0/event/<user_token>")
def endpoint_get_event_data(user_token: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    return success({"event_name": database.hget("event_names", user_token)})


@app.post("/v0/event/<user_token>/check")
def endpoint_check_admin_token(user_token: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    if "admin_token" not in request.json:
        return error(ErrorCode.InternalError)
    admin_token = str(request.json["admin_token"])
    if not is_admin_token_valid(user_token, admin_token):
        return error(ErrorCode.IncorrectAdminToken)
    return success({"event_name": database.hget("event_names", user_token)})


################################################################
#                     ENDPOINTS (IMAGES)                       #
################################################################

def generate_image_thumb(base64_string: str, thumb_width: int, thumb_height: int) -> str:
    image = Image.open(BytesIO(base64.b64decode(base64_string)))
    thumb_image = image.resize(size=(thumb_width, thumb_height))
    buffered = BytesIO()
    thumb_image.save(buffered, format="JPEG")
    thumb_base64 = str(base64.b64encode(buffered.getvalue()))
    return thumb_base64


@app.get("/v0/event/<user_token>/imagecount")
def endpoint_get_image_count(user_token: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    return success(database.llen(image_id_list_name(user_token)))


@app.get("/v0/event/<user_token>/imageids/<first_index>/<last_index>")
def endpoint_get_image_ids(user_token: str, first_index: str, last_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_index = int(first_index)
    last_index = int(last_index)
    if first_index == 0 and last_index == -1:
        return success(database.lrange(image_id_list_name(user_token), first_index, last_index))
    if first_index < 0 or last_index < 0:
        return error(ErrorCode.InternalError)
    return success(database.lrange(image_id_list_name(user_token), first_index, last_index))


@app.post("/v0/event/<user_token>/image")
def endpoint_add_image(user_token: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    if "width" not in request.json:
        return error(ErrorCode.InternalError)
    width = int(request.json["width"])
    if width < 2 or width > 8192:
        return error(ErrorCode.ImageWidthInvalid)

    if "height" not in request.json:
        return error(ErrorCode.InternalError)
    height = int(request.json["height"])
    if height < 2 or height > 8192:
        return error(ErrorCode.ImageHeightInvalid)

    if "description" not in request.json:
        return error(ErrorCode.InternalError)
    description = str(request.json["description"])
    if len(description) > 256:
        return error(ErrorCode.StringTooLong)
    if not filter_user_string(description):
        return error(ErrorCode.StringFailedFiltering)

    if "pixels" not in request.json:
        return error(ErrorCode.InternalError)
    pixels = str(request.json["pixels"])

    album_id: str | None = None
    if "album_id" in request.json:
        album_id = str(request.json["album_id"])
        if album_id not in database.lrange(album_id_list_name(user_token), 0, -1):
            return error(ErrorCode.InvalidAlbumIndex)

    thumb_width = 256 if width > 256 else width
    thumb_height = int(thumb_width * (height / width))
    #thumb_pixels = generate_image_thumb(pixels, thumb_width, thumb_height)
    thumb_pixels = "temp"

    image_time = time.time_ns() // 1000000
    stream_name = image_stream_name(user_token)
    list_name = image_id_list_name(user_token)
    image_id = database.xadd(stream_name, {
        "width": width,
        "height": height,
        "description": description,
        "pixels": pixels,
        "thumbWidth": thumb_width,
        "thumbHeight": thumb_height,
        "thumbPixels": thumb_pixels,
        "time": image_time
    })
    #@TODO: Does this make any sense?
    try:
        transaction = database.pipeline()
        transaction.lpush(list_name, image_id)
        if album_id is not None:
            transaction.lpush(album_image_id_list_name(user_token, album_id), image_id)
        transaction.execute()
    except:
        database.xdel(stream_name, image_id)
        raise
    return success({"image_id": image_id})


@app.delete("/v0/event/<user_token>/image/byid/<image_id>")
def endpoint_delete_image(user_token: str, image_id: str):
    if "admin_token" not in request.json:
        return error(ErrorCode.InternalError)
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    admin_token = str(request.json["admin_token"])
    if not is_admin_token_valid(user_token, admin_token):
        return error(ErrorCode.IncorrectAdminToken)

    current_album_id = None
    for album_id in database.lrange(album_id_list_name(user_token), 0, -1):
        if image_id in database.lrange(album_image_id_list_name(user_token, album_id), 0, -1):
            current_album_id = album_id
            break

    transaction = database.pipeline()
    transaction.ltrim(image_comment_id_list_name(user_token, image_id), 0, -1)
    transaction.xtrim(image_comment_stream_name(user_token, image_id), minid="0-0")
    if current_album_id is not None:
        transaction.lrem(album_image_id_list_name(user_token, current_album_id), 0, image_id)
    transaction.lrem(image_id_list_name(user_token), 0, image_id)
    transaction.xdel(image_stream_name(user_token), image_id)
    transaction.execute()
    return success({})


def get_images_by_ids(*, user_token: str, start_image_id: str, count: int | None, is_thumb: bool):
    width_name = "thumbWidth" if is_thumb else "width"
    height_name = "thumbHeight" if is_thumb else "height"
    pixels_name = "thumbPixels" if is_thumb else "pixels"
    stream = database.xrange(image_stream_name(user_token), start_image_id, "+", count)
    result = []
    for identifier, attributes in stream:
        result.append({
            "image_id": identifier,
            "width": int(attributes[width_name]) if width_name in attributes else 0,
            "height": int(attributes[height_name]) if height_name in attributes else 0,
            "description": attributes["description"],
            "pixels": attributes[pixels_name] if pixels_name in attributes else "",
        })
    return result


@app.get("/v0/event/<user_token>/image/byindex/<image_index>")
def endpoint_get_image_by_index(user_token: str, image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    image_id = database.lindex(image_id_list_name(user_token), int(image_index))
    if image_id is None: return error(ErrorCode.InvalidImageIndex)
    return endpoint_get_image_by_id(user_token, str(image_id))


@app.get("/v0/event/<user_token>/image/byid/<image_id>")
def endpoint_get_image_by_id(user_token: str, image_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    data = get_images_by_ids(user_token=user_token, start_image_id=image_id, count=1, is_thumb=False)
    return success(data)


@app.get("/v0/event/<user_token>/image/byindices/<first_image_index>/<last_image_index>")
def endpoint_get_images_by_indices(user_token: str, first_image_index: str, last_image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_image_index = int(first_image_index)
    last_image_index = int(last_image_index)

    if first_image_index == 0 and last_image_index == -1:
        data = get_images_by_ids(user_token=user_token, start_image_id="0-0", count=None, is_thumb=False)
        return success(data)
    if first_image_index < 0 or last_image_index < 0:
        return error(ErrorCode.InvalidImageIndex)

    result = []
    image_ids = database.lrange(image_id_list_name(user_token), first_image_index, last_image_index)
    for image_id in image_ids:
        data = get_images_by_ids(user_token=user_token, start_image_id=str(image_id), count=1, is_thumb=False)
        result.extend(data)
    return success(result)


################################################################
#                   ENDPOINTS (IMAGE THUMBS)                   #
################################################################

@app.get("/v0/event/<user_token>/imagethumbs/byindex/<image_index>")
def endpoint_get_image_thumb_by_index(user_token: str, image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    image_id = database.lindex(image_id_list_name(user_token), int(image_index))
    if image_id is None: return error(ErrorCode.InvalidImageIndex)
    return endpoint_get_image_thumb_by_id(user_token, str(image_id))


@app.get("/v0/event/<user_token>/imagethumbs/byid/<image_id>")
def endpoint_get_image_thumb_by_id(user_token: str, image_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    data = get_images_by_ids(user_token=user_token, start_image_id=str(image_id), count=1, is_thumb=True)
    return success(data)


@app.get("/v0/event/<user_token>/imagethumbs/byindices/<first_image_index>/<last_image_index>")
def endpoint_get_image_thumbs_by_indices(user_token: str, first_image_index: str, last_image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_image_index = int(first_image_index)
    last_image_index = int(last_image_index)

    if first_image_index == 0 and last_image_index == -1:
        data = get_images_by_ids(user_token=user_token, start_image_id="0-0", count=None, is_thumb=True)
        return success(data)
    if first_image_index < 0 or last_image_index < 0:
        return error(ErrorCode.InvalidImageIndex)

    result = []
    image_ids = database.lrange(image_id_list_name(user_token), first_image_index, last_image_index)
    for image_id in image_ids:
        data = get_images_by_ids(user_token=user_token, start_image_id=str(image_id), count=1, is_thumb=True)
        result.extend(data)
    return success(result)


################################################################
#                    ENDPOINTS (COMMENTS)                      #
################################################################

@app.post("/v0/event/<user_token>/image/byid/<image_id>/comment")
def endpoint_add_comment(user_token: str, image_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    if "text" not in request.json:
        return error(ErrorCode.InternalError)
    text = str(request.json["text"])
    if len(text) > 4096:
        return error(ErrorCode.StringTooLong)
    if not filter_user_string(text):
        return error(ErrorCode.StringFailedFiltering)

    stream_name = image_comment_stream_name(user_token, image_id)
    list_name = image_comment_id_list_name(user_token, image_id)
    comment_time = time.time_ns() // 1000000
    comment_id = database.xadd(stream_name, {
        "text": text,
        "time": comment_time
    })
    try:
        database.lpush(list_name, comment_id)
    except:
        database.xdel(stream_name, comment_id)
        raise
    return success({"comment_id": comment_id, "time": comment_time})


@app.get("/v0/event/<user_token>/image/byid/<image_id>/commentcount")
def endpoint_get_image_comment_count(user_token: str, image_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    return success(database.llen(image_comment_id_list_name(user_token, image_id)))


@app.get("/v0/event/<user_token>/image/byid/<image_id>/commentids/<first_index>/<last_index>")
def endpoint_get_image_comment_ids(user_token: str, image_id: str, first_index: str, last_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_index = int(first_index)
    last_index = int(last_index)
    if (last_index < first_index and last_index != -1) or first_index < 0 or last_index < -1:
        return error(ErrorCode.InternalError)
    return success(database.lrange(image_comment_id_list_name(user_token, image_id), first_index, last_index))


@app.delete("/v0/event/<user_token>/image/byid/<image_id>/comment/byid/<comment_id>")
def endpoint_delete_comment_by_id(user_token: str, image_id: str, comment_id: str):
    if "admin_token" not in request.json:
        return error(ErrorCode.InternalError)
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    admin_token = str(request.json["admin_token"])
    if not is_admin_token_valid(user_token, admin_token):
        return error(ErrorCode.IncorrectAdminToken)

    transaction = database.pipeline()
    transaction.lrem(image_comment_id_list_name(user_token, image_id), 0, comment_id)
    transaction.xdel(image_comment_stream_name(user_token, image_id), comment_id)
    transaction.execute()
    return success({})


def get_image_comment_by_id(user_token: str, image_id: str, comment_id: str):
    stream = database.xrange(image_comment_stream_name(user_token, image_id), comment_id, "+", 1)
    data = None
    for identifier, attributes in stream:
        temp = {"comment_id": identifier}
        temp.update(attributes)
        data = temp
    return data


@app.get("/v0/event/<user_token>/image/byid/<image_id>/comment/byindex/<comment_index>")
def endpoint_get_image_comment_by_index(user_token: str, image_id: str, comment_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    comment_id = database.lindex(image_comment_id_list_name(user_token, image_id), int(comment_index))
    if comment_id is None:
        return error(ErrorCode.InvalidCommentIndex)
    return endpoint_get_image_comment_by_id(user_token, image_id, str(comment_id))


@app.get("/v0/event/<user_token>/image/byid/<image_id>/comment/byid/<comment_id>")
def endpoint_get_image_comment_by_id(user_token: str, image_id: str, comment_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    data = get_image_comment_by_id(user_token, image_id, comment_id)
    return success([data] if data is not None else [])


#@TODO: This looks bad.
@app.get("/v0/event/<user_token>/image/byid/<image_id>/comment/byindices/<first_comment_index>/<last_comment_index>")
def endpoint_get_image_comments_by_indices(user_token: str, image_id: str, first_comment_index: str,
                                           last_comment_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    first_comment_index = int(first_comment_index)
    last_comment_index = int(last_comment_index)
    if first_comment_index == 0 and last_comment_index == -1:
        stream = database.xrange(image_comment_stream_name(user_token, image_id), "0-0", "+")
        comment_datas = []
        for identifier, attributes in stream:
            temp = {"comment_id": identifier}
            temp.update(attributes)
            comment_datas.append(temp)
        return success(comment_datas)
    else:
        if first_comment_index < 0 or last_comment_index < 0 or last_comment_index < first_comment_index:
            return error(ErrorCode.InvalidCommentIndex)

        image_comment_id_list = image_comment_id_list_name(user_token, image_id)
        comment_datas = []
        for comment_index in range(first_comment_index, last_comment_index + 1):
            comment_id = database.lindex(image_comment_id_list, comment_index)
            if comment_id is None: continue
            data = get_image_comment_by_id(user_token, image_id, str(comment_id))
            if data is None: continue
            comment_datas.append(data)
        return success(comment_datas)


################################################################
#                     ENDPOINTS (ALBUMS)                       #
################################################################

@app.post("/v0/event/<user_token>/album")
def endpoint_create_album(user_token: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    if "name" not in request.json:
        return error(ErrorCode.InternalError)
    album_name = str(request.json["name"])
    if len(album_name) > 512:
        return error(ErrorCode.StringTooLong)
    if not filter_user_string(album_name):
        return error(ErrorCode.StringFailedFiltering)

    album_time = time.time_ns() // 1000000
    stream_name = album_stream_name(user_token)
    list_name = album_id_list_name(user_token)
    album_id = database.xadd(stream_name, {
        "name": album_name,
        "time": album_time
    })
    #@TODO: Does this make any sense?
    try:
        database.lpush(list_name, album_id)
    except:
        database.xdel(stream_name, album_id)
        raise
    return success({"album_id": album_id, "time": album_time})


@app.delete("/v0/event/<user_token>/album/byid/<album_id>")
def endpoint_delete_album_by_id(user_token: str, album_id: str):
    if "admin_token" not in request.json:
        return error(ErrorCode.InternalError)
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    admin_token = str(request.json["admin_token"])
    if not is_admin_token_valid(user_token, admin_token):
        return error(ErrorCode.IncorrectAdminToken)

    transaction = database.pipeline()
    transaction.ltrim(album_image_id_list_name(user_token, album_id), 0, -1)
    transaction.lrem(album_id_list_name(user_token), 0, album_id)
    transaction.xdel(album_stream_name(user_token), album_id)
    transaction.execute()
    return success({})


def get_albums_by_ids(*, user_token: str, start_album_id: str, count: int | None):
    stream = database.xrange(album_stream_name(user_token), start_album_id, "+", count)
    result = []
    for identifier, attributes in stream:
        image_count = database.llen(album_image_id_list_name(user_token, identifier))
        if image_count is None: image_count = 0
        result.append({
            "album_id": identifier,
            "name": attributes["name"],
            "image_count": int(image_count),
            "time": int(attributes["time"])
        })
    return result


@app.get("/v0/event/<user_token>/album/byindices/<first_album_index>/<last_album_index>")
def endpoint_get_albums_by_indices(user_token: str, first_album_index: str, last_album_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_album_index = int(first_album_index)
    last_album_index = int(last_album_index)

    if first_album_index == 0 and last_album_index == -1:
        data = get_albums_by_ids(user_token=user_token, start_album_id="0-0", count=None)
        return success(data)
    if first_album_index < 0 or last_album_index < 0:
        return error(ErrorCode.InvalidAlbumIndex)

    result = []
    album_ids = database.lrange(album_id_list_name(user_token), first_album_index, last_album_index)
    for album_id in album_ids:
        data = get_albums_by_ids(user_token=user_token, start_album_id=str(album_id), count=1)
        result.extend(data)
    return success(result)


@app.get("/v0/event/<user_token>/album/byindex/<album_index>")
def endpoint_get_album_by_index(user_token: str, album_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    album_id = database.lindex(album_id_list_name(user_token), int(album_index))
    if album_id is None: return error(ErrorCode.InvalidAlbumIndex)
    return endpoint_get_album_by_id(user_token, str(album_id))


@app.get("/v0/event/<user_token>/album/byid/<album_id>")
def endpoint_get_album_by_id(user_token: str, album_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    data = get_albums_by_ids(user_token=user_token, start_album_id=str(album_id), count=1)
    return success(data)


@app.get("/v0/event/<user_token>/album/byid/<album_id>/imageids/<first_image_index>/<last_image_index>")
def endpoint_get_album_image_ids(user_token: str, album_id: str, first_image_index: str, last_image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    first_image_index = int(first_image_index)
    last_image_index = int(last_image_index)

    if album_id not in database.lrange(album_id_list_name(user_token), 0, -1):
        return error(ErrorCode.IncorrectAlbumId)

    if first_image_index == 0 and last_image_index == -1:
        result = database.lrange(album_image_id_list_name(user_token, album_id), 0, -1)
        return success(result)
    if first_image_index < 0 or last_image_index < 0:
        return error(ErrorCode.InvalidImageIndex)

    result = database.lrange(album_image_id_list_name(user_token, album_id), first_image_index, last_image_index)
    return success(result)


@app.get("/v0/event/<user_token>/album/byid/<album_id>/imagecount")
def endpoint_get_album_image_count(user_token: str, album_id: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    if album_id not in database.lrange(album_id_list_name(user_token), 0, -1):
        return error(ErrorCode.IncorrectAlbumId)
    length = database.llen(album_image_id_list_name(user_token, album_id))
    if length is None:
        return error(ErrorCode.InternalError)
    return success(int(length))


@app.get("/v0/event/<user_token>/album/byid/<album_id>/image/byindex/<image_index>")
def endpoint_get_album_image_by_index(user_token: str, album_id: str, image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    image_index = int(image_index)
    if image_index < 0:
        return error(ErrorCode.InvalidImageIndex)
    image_id = database.lindex(album_image_id_list_name(user_token, album_id), image_index)
    if image_id is None:
        return error(ErrorCode.InvalidImageIndex)
    return endpoint_get_image_by_id(user_token, image_id)


@app.get("/v0/event/<user_token>/album/byid/<album_id>/imagethumbs/byindex/<image_index>")
def endpoint_get_album_image_thumb_by_index(user_token: str, album_id: str, image_index: str):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)
    image_index = int(image_index)
    if image_index < 0:
        return error(ErrorCode.InvalidImageIndex)
    image_id = database.lindex(album_image_id_list_name(user_token, album_id), image_index)
    if image_id is None:
        return error(ErrorCode.InvalidImageIndex)
    return endpoint_get_image_thumb_by_id(user_token, image_id)


@app.get("/v0/event/<user_token>/PDF")
def endpoint_create_pdf_album(user_token):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    images_list = get_images_by_ids(user_token=user_token, start_image_id="-", count=None, is_thumb=False)
    if len(images_list) == 0:
        return error(ErrorCode.InternalError)

    album_title = "wedding_album"
    images_PIL = images_b64_to_pillow([(image["pixels"], image["width"], image["height"]) for image in images_list])

    pdf_buffer = PDF_from_pillow(images_PIL, album_title)
    return Response(
        pdf_buffer,
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{album_title}.pdf"'
        }
    )


def images_b64_to_pillow(image_list):
    images_PIL = []
    for b64_image, width, height in image_list:
        image = Image.open(BytesIO(base64.b64decode(b64_image)))
        images_PIL.append(image)
    return images_PIL


def draw_page_number(c, page_num):
    c.setFont("Helvetica", 10)
    c.drawString(inch, 0.75 * inch, f"{page_num}")


def PDF_from_pillow(images_PIL, album_title):
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    page_width, page_height = letter

    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(page_width / 2.0, page_height / 2.0, album_title)
    c.showPage()

    for i, image in enumerate(images_PIL):
        # c.drawImage(image=image, x=inch, y=inch, width=page_width - 2 * inch, height=page_height - 2 * inch,
        #             preserveAspectRatio=True)
        c.drawInlineImage(image=image, x=inch, y=inch, width=page_width - 2 * inch, height=page_height - 2 * inch)
        draw_page_number(c, i + 1)
        c.showPage()

    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer


@app.get("/v0/event/<user_token>/DOCX")
def endpoint_create_docx_album(user_token):
    if not does_event_exist(user_token):
        return error(ErrorCode.IncorrectUserToken)

    images_list = get_images_by_ids(user_token=user_token, start_image_id="-", count=None, is_thumb=False)
    if len(images_list) == 0:
        return error(ErrorCode.InternalError)


    album_title = "Wedding album"
    images_PIL = images_b64_to_pillow([(image["pixels"], image["width"], image["height"]) for image in images_list])

    docx_buffer = DOCX_from_pillow(images_PIL, album_title)
    return Response(
        docx_buffer,
        mimetype='application/docx',
        headers={
            'Content-Disposition': f'attachment; filename="{album_title}.docx'
        }
    )


def add_image(doc, image_PIL, max_width=6.5, max_height=9.0):
    width, height = image_PIL.size
    aspect_ratio = width / height
    if width / max_width > height / max_height:
        new_width = max_width
        new_height = max_width / aspect_ratio
    else:
        new_height = max_height
        new_width = max_height * aspect_ratio
    image_stream = BytesIO()
    image_PIL.save(image_stream, format="JPEG")
    image_stream.seek(0)

    doc.add_picture(image_stream, width=Inches(new_width), height=Inches(new_height))


def DOCX_from_pillow(images_PIL, title):
    docx_buffer = BytesIO()

    doc = docx.Document()
    styles = doc.styles
    styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.size = docx.shared.Pt(30)
    p = doc.add_heading(title, level=2)
    p.alignment = 1
    for i, image_PIL in enumerate(images_PIL):
        if i > 0:
            doc.add_page_break()
        add_image(doc, image_PIL)

    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer
